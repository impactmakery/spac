"""Demo dataset: two municipalities with real content, for pilots and handover.

Usage:
    python scripts/seed_demo.py            # create (skips if already present)
    python scripts/seed_demo.py --index    # also run the ingestion worker inline

Idempotent: it keys on the demo municipality names and refuses to run twice, so
it is safe against a database that already has real data in it.

There is no un-seed. Removing this data by hand means FK-ordered deletes across
roughly ten tables plus the R2 objects — ask for a `--remove` flag rather than
improvising it against production.

The content is deliberately chosen to demonstrate the permission model rather
than to look pretty. Karmiel's education department holds a staffing figure
that appears nowhere else, so asking about it as a Nahariya user must return
"not covered" — that is the acceptance item from the scope appendix, visible in
the product rather than only in the test suite.
"""

import argparse
import io
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from sqlalchemy import create_engine, func, select  # noqa: E402
from sqlalchemy.orm import Session, sessionmaker  # noqa: E402

from app.core.config import get_settings  # noqa: E402
from app.core.security import hash_password  # noqa: E402
from app.models import (  # noqa: E402
    BoardComment,
    BoardItem,
    BoardLike,
    Category,
    Department,
    DepartmentFile,
    DepartmentPost,
    KbDocument,
    Municipality,
    User,
)
from app.services.ingestion import enqueue, run_pending_jobs  # noqa: E402
from app.services.storage import get_storage  # noqa: E402

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PASSWORD = "demo-password-1"  # noqa: S105 — demo credentials, documented in the README

CATEGORIES = [
    ("חינוך", "Education"),
    ("רווחה", "Welfare"),
    ("תכנון ובנייה", "Planning and building"),
    ("איכות הסביבה", "Environment"),
    ("תקציב ורכש", "Budget and procurement"),
]

# (title, filename, paragraphs)
KB_DOCS = [
    (
        "נוהל טיפול בפניות תושבים",
        "resident-enquiries.docx",
        [
            "נוהל זה מסדיר את הטיפול בפניות תושבים בכלל הרשויות המשתתפות.",
            "פנייה שהתקבלה במוקד העירוני תיענה תוך שלושה ימי עבודה. פנייה דחופה "
            "בנושא בטיחות או תשתיות תטופל תוך 24 שעות.",
            "כל פנייה מתועדת במערכת ומקבלת מספר מעקב. התושב מקבל הודעת עדכון בסיום "
            "הטיפול, גם אם הפנייה נדחתה.",
            "פנייה שאינה בסמכות הרשות תועבר לגורם המוסמך תוך חמישה ימי עבודה, "
            "והתושב יעודכן על ההעברה.",
        ],
    ),
    (
        "מדריך היערכות לחירום ברשות המקומית",
        "emergency-preparedness.docx",
        [
            "המדריך מפרט את חלוקת האחריות ברשות המקומית בשעת חירום.",
            "מנהל אגף הביטחון מרכז את ההיערכות. כל אגף ממנה נציג חירום ומעדכן את "
            "פרטי הקשר שלו אחת לרבעון.",
            "מרכזי הקליטה נפתחים בהוראת ראש הרשות. רשימת המבנים המיועדים מתעדכנת "
            "בתחילת כל שנה קלנדרית.",
            "תרגיל חירום עירוני מתקיים פעמיים בשנה, במרץ ובאוקטובר.",
        ],
    ),
    (
        "Municipal budget planning guide 2026",
        "budget-guide-2026.docx",
        [
            "This guide covers the 2026 budget planning cycle for participating "
            "municipalities.",
            "The budget planning cycle begins in November. Departments submit "
            "their draft requests by 15 December, and the finance committee "
            "reviews them in January.",
            "Procurement above 150,000 NIS requires a tender under regulation "
            "17.3 of the municipal procurement code. Applications use form 4B.",
            "Mid-year adjustments are permitted once, in July, and may not "
            "exceed 8 percent of the approved departmental budget.",
        ],
    ),
]


def _docx(paragraphs: list[str]) -> bytes:
    import docx

    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _user(db: Session, email: str, name: str, role: str, **kw) -> User:
    existing = db.scalar(select(User).where(func.lower(User.email) == email.lower()))
    if existing:
        return existing
    u = User(
        email=email,
        name=name,
        role=role,
        status="active",
        password_hash=hash_password(PASSWORD),
        **kw,
    )
    db.add(u)
    return u


def seed(db: Session, *, index: bool) -> None:
    if db.scalar(select(Municipality).where(Municipality.name == "עיריית נהריה")):
        print("demo data already present — nothing to do")
        return

    storage = get_storage()
    stamp = int(time.time())

    categories = {}
    for he, en in CATEGORIES:
        cat = db.scalar(select(Category).where(Category.name_he == he))
        if cat is None:
            cat = Category(name_he=he, name_en=en)
            db.add(cat)
        categories[en] = cat

    nahariya = Municipality(name="עיריית נהריה")
    karmiel = Municipality(name="עיריית כרמיאל")
    db.add_all([nahariya, karmiel])

    n_welfare = Department(municipality=nahariya, name="רווחה")
    n_education = Department(municipality=nahariya, name="חינוך")
    n_sanitation = Department(municipality=nahariya, name="תברואה")
    k_welfare = Department(municipality=karmiel, name="רווחה")
    k_education = Department(municipality=karmiel, name="חינוך")
    db.add_all([n_welfare, n_education, n_sanitation, k_welfare, k_education])
    db.flush()

    admin = _user(db, "admin@tomorrow-hub.org", "מנהל מערכת", "system_admin", language="he")
    n_admin = _user(
        db, "admin.nahariya@tomorrow-hub.org", "מנהלת נהריה", "municipality_admin",
        municipality_id=nahariya.id, language="he",
    )
    k_admin = _user(
        db, "admin.karmiel@tomorrow-hub.org", "מנהל כרמיאל", "municipality_admin",
        municipality_id=karmiel.id, language="he",
    )
    n_welfare_user = _user(
        db, "welfare.nahariya@tomorrow-hub.org", "עובדת רווחה נהריה", "department_user",
        municipality_id=nahariya.id, language="he",
    )
    n_education_user = _user(
        db, "education.nahariya@tomorrow-hub.org", "רכז חינוך נהריה", "department_user",
        municipality_id=nahariya.id, language="he",
    )
    k_education_user = _user(
        db, "education.karmiel@tomorrow-hub.org", "רכזת חינוך כרמיאל", "department_user",
        municipality_id=karmiel.id, language="he",
    )
    db.flush()
    n_welfare_user.departments.append(n_welfare)
    n_education_user.departments.append(n_education)
    k_education_user.departments.append(k_education)

    # --- global knowledge base -------------------------------------------------
    for title, filename, paragraphs in KB_DOCS:
        content = _docx(paragraphs)
        doc = KbDocument(
            title=title, filename=filename, storage_key="", size_bytes=len(content),
            content_type=DOCX_MIME, uploader_id=admin.id,
        )
        db.add(doc)
        db.flush()
        doc.storage_key = f"kb/{doc.id}/{stamp}/{filename}"
        storage.put(doc.storage_key, content, DOCX_MIME)
        enqueue(
            db, source_type="kb", source_id=doc.id, visibility="global",
            storage_key=doc.storage_key, ext="docx", title=doc.title,
        )

    # --- boards ---------------------------------------------------------------
    global_item = BoardItem(
        title="כנס מנהלי רווחה ארצי 2026",
        description="הכנס יתקיים ב-12 במרץ 2026 במרכז הכנסים בתל אביב. "
        "ההרשמה נסגרת שבועיים לפני המועד. הנסיעה מתוקצבת על ידי הרשות השולחת.",
        category_id=categories["Welfare"].id, scope="global", author_id=admin.id,
        link_url="https://example.org/welfare-conference-2026", indexing_status="pending",
    )
    nahariya_item = BoardItem(
        title="שעות פתיחה מעודכנות של מרכז המיחזור",
        description="מרכז המיחזור העירוני פתוח בימים א'-ה' בין 08:00 ל-16:00, "
        "וביום שישי עד 12:00. איסוף גזם מתבצע בימי שני ורביעי.",
        category_id=categories["Environment"].id, scope="municipality",
        municipality_id=nahariya.id, author_id=n_admin.id, indexing_status="pending",
    )
    karmiel_item = BoardItem(
        title="עדכון נוהל רכש לשנת 2026",
        description="רכש מעל 150,000 ש\"ח מחייב מכרז. הטפסים המעודכנים נמצאים "
        "במערכת הרכש העירונית.",
        category_id=categories["Budget and procurement"].id, scope="municipality",
        municipality_id=karmiel.id, author_id=k_admin.id, indexing_status="pending",
    )
    db.add_all([global_item, nahariya_item, karmiel_item])
    db.flush()

    db.add_all([
        BoardComment(item_id=global_item.id, author_id=n_welfare_user.id,
                     body="האם יש אפשרות להשתתפות מקוונת?"),
        BoardComment(item_id=global_item.id, author_id=admin.id,
                     body="כן, קישור לשידור יישלח לנרשמים יומיים לפני הכנס."),
        BoardLike(item_id=global_item.id, user_id=n_welfare_user.id),
        BoardLike(item_id=global_item.id, user_id=k_education_user.id),
        BoardLike(item_id=nahariya_item.id, user_id=n_education_user.id),
    ])

    for item, visibility, muni in (
        (global_item, "global", None),
        (nahariya_item, "municipality", nahariya),
        (karmiel_item, "municipality", karmiel),
    ):
        enqueue(
            db, source_type="board", source_id=item.id, visibility=visibility,
            text_content=item.description or "", title=item.title,
            municipality_id=muni.id if muni else None,
        )

    # --- department content ---------------------------------------------------
    # The figure below exists only here. A Nahariya user asking about it must get
    # the "not covered" reply — the demo of the permission boundary.
    k_edu_content = _docx([
        "תוכנית כוח האדם של אגף החינוך בכרמיאל לשנת 2026.",
        "התוכנית מקצה ארבעה תקנים חדשים: שני יועצים חינוכיים, מדריך הכלה אחד "
        "ורכז נוער אחד.",
        "התקנים ייכנסו לתוקף בספטמבר 2026 בכפוף לאישור התקציב.",
    ])
    k_edu_file = DepartmentFile(
        department_id=k_education.id, uploader_id=k_admin.id,
        filename="staffing-plan-2026.docx", storage_key="",
        size_bytes=len(k_edu_content), content_type=DOCX_MIME,
    )
    db.add(k_edu_file)
    db.flush()
    k_edu_file.storage_key = f"department/{k_education.id}/{stamp}/{k_edu_file.filename}"
    storage.put(k_edu_file.storage_key, k_edu_content, DOCX_MIME)
    enqueue(
        db, source_type="department", source_id=k_edu_file.id, visibility="department",
        storage_key=k_edu_file.storage_key, ext="docx", title=k_edu_file.filename,
        municipality_id=karmiel.id, department_id=k_education.id,
    )

    n_welfare_content = _docx([
        "נוהל קליטת פניות באגף הרווחה בנהריה.",
        "פנייה חדשה נקלטת על ידי עובד התורנות ומועברת לעובד סוציאלי תוך יומיים.",
        "מקרה דחוף מועבר מיידית לראש הצוות, ללא המתנה לתורנות.",
    ])
    n_welfare_file = DepartmentFile(
        department_id=n_welfare.id, uploader_id=n_admin.id,
        filename="welfare-intake.docx", storage_key="",
        size_bytes=len(n_welfare_content), content_type=DOCX_MIME,
    )
    db.add(n_welfare_file)
    db.flush()
    n_welfare_file.storage_key = (
        f"department/{n_welfare.id}/{stamp}/{n_welfare_file.filename}"
    )
    storage.put(n_welfare_file.storage_key, n_welfare_content, DOCX_MIME)
    enqueue(
        db, source_type="department", source_id=n_welfare_file.id,
        visibility="department", storage_key=n_welfare_file.storage_key, ext="docx",
        title=n_welfare_file.filename, municipality_id=nahariya.id,
        department_id=n_welfare.id,
    )

    post = DepartmentPost(
        department_id=n_welfare.id, author_id=n_welfare_user.id,
        body="תזכורת: ישיבת צוות רווחה מתקיימת כל יום שלישי בשעה 09:00 בחדר הישיבות "
        "בקומה השנייה. נא להביא את סיכומי המקרים הפתוחים.",
    )
    db.add(post)
    db.flush()
    enqueue(
        db, source_type="department", source_id=post.id, visibility="department",
        text_content=post.body, municipality_id=nahariya.id, department_id=n_welfare.id,
    )

    db.commit()
    print("seeded 2 municipalities, 5 departments, 6 users, 3 KB documents,")
    print("3 board items, 2 department files, 1 department post")

    if index:
        processed = 0
        while True:
            n = run_pending_jobs(db, limit=20)
            processed += n
            if n == 0:
                break
        print(f"indexed {processed} source(s)")
    else:
        print("run the ingestion worker (or pass --index) to embed the content")

    print(f"\nall demo accounts use the password: {PASSWORD}")
    print("  admin@tomorrow-hub.org             system admin")
    print("  admin.nahariya@tomorrow-hub.org    municipality admin (Nahariya)")
    print("  admin.karmiel@tomorrow-hub.org     municipality admin (Karmiel)")
    print("  welfare.nahariya@tomorrow-hub.org  department user (Nahariya / Welfare)")
    print("  education.nahariya@tomorrow-hub.org department user (Nahariya / Education)")
    print("  education.karmiel@tomorrow-hub.org department user (Karmiel / Education)")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--index", action="store_true",
        help="run the ingestion worker inline so the content is searchable immediately",
    )
    args = parser.parse_args()

    engine = create_engine(get_settings().database_url)
    with sessionmaker(bind=engine)() as db:
        seed(db, index=args.index)


if __name__ == "__main__":
    main()
