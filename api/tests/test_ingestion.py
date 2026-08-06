import io

import pytest


@pytest.fixture()
def files_dir(tmp_path, monkeypatch):
    from app.core.config import get_settings

    monkeypatch.setattr(get_settings(), "files_dir", str(tmp_path))
    return tmp_path


def _docx_bytes(text) -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_doc(db, storage_key="kb/test/doc.docx"):
    from app.models import KbDocument

    doc = KbDocument(
        title="Guidelines",
        filename="doc.docx",
        storage_key=storage_key,
        size_bytes=1234,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    db.add(doc)
    db.commit()
    return doc


def test_ingestion_happy_path(db, files_dir):
    from app.models import Chunk
    from app.services.ingestion import enqueue, run_pending_jobs
    from app.services.storage import get_storage

    doc = _make_doc(db)
    get_storage().put(doc.storage_key, _docx_bytes("תקציב הרווחה לשנת 2026."), "x")
    enqueue(
        db, source_type="kb", source_id=doc.id, visibility="global",
        storage_key=doc.storage_key, ext="docx",
    )
    db.commit()

    assert run_pending_jobs(db) == 1
    db.expire_all()
    assert doc.status == "indexed"
    chunks = db.query(Chunk).filter(Chunk.source_id == doc.id).all()
    assert len(chunks) == 1
    assert "הרווחה" in chunks[0].content
    assert chunks[0].visibility == "global" and chunks[0].source_type == "kb"
    assert len(chunks[0].embedding) == 1536


def test_ingestion_retries_then_not_indexable(db, files_dir):
    from datetime import UTC, datetime

    from app.models import IngestionJob
    from app.services.ingestion import enqueue, run_pending_jobs
    from app.services.storage import get_storage

    doc = _make_doc(db, storage_key="kb/test/corrupt.docx")
    get_storage().put(doc.storage_key, b"this is not a docx", "x")
    enqueue(
        db, source_type="kb", source_id=doc.id, visibility="global",
        storage_key=doc.storage_key, ext="docx",
    )
    db.commit()

    for attempt in (1, 2, 3):
        # make any backoff due immediately
        db.query(IngestionJob).update({"run_after": datetime.now(UTC)})
        db.commit()
        assert run_pending_jobs(db) == 1
        db.expire_all()
        job = db.query(IngestionJob).one()
        if attempt < 3:
            assert job.status == "queued" and job.attempts == attempt
            assert doc.status == "processing"
        else:
            assert job.status == "failed" and job.attempts == 3
            assert doc.status == "not_indexable"
            assert job.last_error


def test_reingest_replaces_chunks(db, files_dir):
    from app.models import Chunk
    from app.services.ingestion import enqueue, run_pending_jobs
    from app.services.storage import get_storage

    doc = _make_doc(db)
    get_storage().put(doc.storage_key, _docx_bytes("First version."), "x")
    enqueue(db, source_type="kb", source_id=doc.id, visibility="global",
            storage_key=doc.storage_key, ext="docx")
    db.commit()
    run_pending_jobs(db)

    get_storage().put(doc.storage_key, _docx_bytes("Second version entirely."), "x")
    enqueue(db, source_type="kb", source_id=doc.id, visibility="global",
            storage_key=doc.storage_key, ext="docx")
    db.commit()
    run_pending_jobs(db)

    chunks = db.query(Chunk).filter(Chunk.source_id == doc.id).all()
    assert len(chunks) == 1
    assert "Second version" in chunks[0].content


def test_source_deleted_mid_run_does_not_kill_the_worker(db, files_dir, monkeypatch):
    """Deleting a document removes its job rows; the worker must survive that
    rather than crash the whole cycle on a vanished row."""
    from sqlalchemy import delete as sql_delete

    from app.models import IngestionJob, KbDocument
    from app.services import ingestion
    from app.services.ingestion import enqueue, run_pending_jobs
    from app.services.storage import get_storage

    doc = _make_doc(db)
    get_storage().put(doc.storage_key, _docx_bytes("content"), "x")
    enqueue(db, source_type="kb", source_id=doc.id, visibility="global",
            storage_key=doc.storage_key, ext="docx")
    db.commit()

    real_process = ingestion._process

    def delete_then_process(session, job):
        # simulate the API deleting the document while this job is in flight
        session.execute(sql_delete(IngestionJob).where(IngestionJob.id == job.id))
        session.execute(sql_delete(KbDocument).where(KbDocument.id == doc.id))
        session.commit()
        return real_process(session, job)

    monkeypatch.setattr(ingestion, "_process", delete_then_process)

    # must not raise, and must leave nothing behind
    run_pending_jobs(db)
    db.expire_all()
    assert db.query(IngestionJob).count() == 0
    assert db.query(KbDocument).count() == 0


def test_document_title_lands_on_every_chunk(db, files_dir):
    """A passage cited from the middle of a long file must still name its
    source, and the title sharpens the embedding of every chunk."""
    from app.models import Chunk
    from app.rag.chunking import CHUNK_TOKENS
    from app.services.ingestion import enqueue, run_pending_jobs
    from app.services.storage import get_storage

    doc = _make_doc(db, storage_key="kb/test/long.docx")
    body = "\n\n".join(
        f"Section {i}. " + " ".join(f"detail{j}" for j in range(200))
        for i in range(12)
    )
    get_storage().put(doc.storage_key, _docx_bytes(body), "x")
    enqueue(
        db, source_type="kb", source_id=doc.id, visibility="global",
        storage_key=doc.storage_key, ext="docx", title=doc.title,
    )
    db.commit()
    run_pending_jobs(db)

    chunks = db.query(Chunk).filter(Chunk.source_id == doc.id).all()
    assert len(chunks) > 1, "the fixture body must be long enough to split"
    assert all(c.content.startswith("Guidelines") for c in chunks)
    from app.rag.chunking import token_len

    assert all(token_len(c.content) <= CHUNK_TOKENS for c in chunks)


def test_titled_source_with_no_body_still_indexes(db, files_dir):
    """A board link item has a title and no file; it must stay searchable."""
    from app.models import Chunk
    from app.services.ingestion import enqueue, run_pending_jobs

    doc = _make_doc(db, storage_key="kb/test/empty.docx")
    enqueue(
        db, source_type="board", source_id=doc.id, visibility="global",
        text_content="", title="Recycling centre opening hours",
    )
    db.commit()
    run_pending_jobs(db)

    chunks = db.query(Chunk).filter(Chunk.source_id == doc.id).all()
    assert [c.content for c in chunks] == ["Recycling centre opening hours"]
