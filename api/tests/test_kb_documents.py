import io

import pytest


@pytest.fixture()
def files_dir(tmp_path, monkeypatch):
    from app.core.config import get_settings

    monkeypatch.setattr(get_settings(), "files_dir", str(tmp_path))
    return tmp_path


@pytest.fixture()
def world(db):
    from app.core.security import hash_password
    from app.models import Municipality, User

    m1 = Municipality(name="City One")
    pw = hash_password("kb-password-111")
    sysadmin = User(email="sys@x.org", role="system_admin", status="active",
                    password_hash=pw, name="Root")
    a1 = User(email="a1@x.org", role="municipality_admin", municipality=m1,
              status="active", password_hash=pw, name="Admin")
    u1 = User(email="u1@x.org", role="department_user", municipality=m1,
              status="active", password_hash=pw, name="Worker")
    db.add_all([m1, sysadmin, a1, u1])
    db.commit()
    return {"m1": m1, "sys": sysadmin, "a1": a1, "u1": u1}


def auth(client, email):
    r = client.post("/api/auth/login", json={"email": email, "password": "kb-password-111"})
    assert r.status_code == 200
    return {"Authorization": f"Bearer {r.json()['access_token']}"}


def _docx(text="Municipal waste collection guidelines 2026") -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph(text)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _upload(client, headers, filename="guide.docx", content=None, title=None):
    data = {"title": title} if title else {}
    return client.post(
        "/api/kb-documents",
        files={"file": (filename, content or _docx(), DOCX_MIME)},
        data=data,
        headers=headers,
    )


def test_upload_list_download_flow(client, db, world, files_dir):
    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "a1@x.org")
    r = _upload(client, headers, title="Waste Guide")
    assert r.status_code == 201, r.text
    doc_id = r.json()["id"]
    assert r.json()["status"] == "pending"
    assert r.json()["municipality_name"] == "City One"

    run_pending_jobs(db)

    # the library is browsable by administrators
    rows = client.get("/api/kb-documents", headers=headers).json()
    assert [d["title"] for d in rows] == ["Waste Guide"]
    assert rows[0]["status"] == "indexed"

    # a department user may still open the document a citation points at,
    # or the assistant's sources would be uncheckable
    user_headers = auth(client, "u1@x.org")
    detail = client.get(f"/api/kb-documents/{doc_id}", headers=user_headers).json()
    assert detail["download_url"].startswith("/api/files/kb/")
    # signed URL serves without auth
    assert client.get(detail["download_url"]).status_code == 200

    assert db.query(Chunk).count() == 1

    # search
    assert client.get("/api/kb-documents?search=waste", headers=headers).json()
    assert client.get("/api/kb-documents?search=zzz", headers=headers).json() == []


def test_department_user_cannot_upload(client, world, files_dir):
    headers = auth(client, "u1@x.org")
    assert _upload(client, headers).status_code == 404


def test_upload_rejects_bad_files(client, world, files_dir):
    headers = auth(client, "a1@x.org")
    r = client.post(
        "/api/kb-documents",
        files={"file": ("run.exe", b"MZ...", "application/octet-stream")},
        headers=headers,
    )
    assert r.status_code == 415


def test_replace_keeps_id_and_reindexes(client, db, world, files_dir):
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "a1@x.org")
    doc_id = _upload(client, headers).json()["id"]
    run_pending_jobs(db)

    r = client.post(
        f"/api/kb-documents/{doc_id}/replace",
        files={"file": ("v2.docx", _docx("Second edition of the rules"), DOCX_MIME)},
        headers=headers,
    )
    assert r.status_code == 200
    assert r.json()["id"] == doc_id and r.json()["status"] == "pending"
    run_pending_jobs(db)

    from app.models import Chunk

    chunks = db.query(Chunk).all()
    assert len(chunks) == 1 and "Second edition" in chunks[0].content


def test_replace_delete_scoped_to_uploader_or_sysadmin(client, db, world, files_dir):
    from app.core.security import hash_password
    from app.models import Municipality, User

    m2 = Municipality(name="City Two")
    a2 = User(email="a2@x.org", role="municipality_admin", municipality=m2,
              status="active", password_hash=hash_password("kb-password-111"))
    db.add_all([m2, a2])
    db.commit()

    doc_id = _upload(client, auth(client, "a1@x.org")).json()["id"]
    other = auth(client, "a2@x.org")
    assert (
        client.post(
            f"/api/kb-documents/{doc_id}/replace",
            files={"file": ("x.docx", _docx(), DOCX_MIME)},
            headers=other,
        ).status_code
        == 404
    )
    assert client.delete(f"/api/kb-documents/{doc_id}", headers=other).status_code == 404
    # sysadmin can delete anything
    assert (
        client.delete(f"/api/kb-documents/{doc_id}", headers=auth(client, "sys@x.org"))
        .status_code
        == 200
    )


def test_delete_cascades_chunks_same_transaction(client, db, world, files_dir):
    from app.models import Chunk, IngestionJob
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "a1@x.org")
    doc_id = _upload(client, headers).json()["id"]
    run_pending_jobs(db)
    assert db.query(Chunk).count() == 1

    assert client.delete(f"/api/kb-documents/{doc_id}", headers=headers).status_code == 200
    db.expire_all()
    assert db.query(Chunk).count() == 0
    assert db.query(IngestionJob).count() == 0
    assert client.get(f"/api/kb-documents/{doc_id}", headers=headers).status_code == 404


# --- text preview -------------------------------------------------------------
# A browser cannot render Word, PowerPoint or Excel in a frame, and the
# knowledge base is mostly Word. The text is already extracted for the
# assistant, so previewing it costs nothing new.


def test_word_document_text_can_be_previewed(client, world, files_dir):
    sys_headers = auth(client, "sys@x.org")
    doc_id = _upload(
        client, sys_headers, filename="procedure.docx",
        content=_docx("Waste collection operates on Mondays and Thursdays"),
    ).json()["id"]

    res = client.get(f"/api/kb-documents/{doc_id}/text", headers=sys_headers)
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["available"] is True
    assert "Mondays and Thursdays" in body["text"]
    assert body["truncated"] is False


def test_a_long_document_is_truncated_rather_than_sent_whole(
    client, world, files_dir, monkeypatch
):
    sys_headers = auth(client, "sys@x.org")
    from app.routers import kb_documents

    monkeypatch.setattr(kb_documents, "MAX_PREVIEW_CHARS", 50)
    doc_id = _upload(
        client, sys_headers, filename="long.docx",
        content=_docx("word " * 500),
    ).json()["id"]

    body = client.get(f"/api/kb-documents/{doc_id}/text", headers=sys_headers).json()
    assert body["truncated"] is True
    assert len(body["text"]) == 50


def test_a_format_with_no_extractor_reports_no_preview(
    client, world, files_dir, db
):
    """Rather than erroring: the page falls back to 'download the file'."""
    sys_headers = auth(client, "sys@x.org")
    from app.models import KbDocument

    doc = KbDocument(
        title="Archive", filename="bundle.zip", storage_key="kb/x/bundle.zip",
        size_bytes=10, content_type="application/octet-stream",
    )
    db.add(doc)
    db.commit()

    body = client.get(f"/api/kb-documents/{doc.id}/text", headers=sys_headers).json()
    assert body["available"] is False and body["text"] == ""


def test_a_missing_file_does_not_break_the_page(client, world, files_dir, db):
    sys_headers = auth(client, "sys@x.org")
    from app.models import KbDocument

    doc = KbDocument(
        title="Gone", filename="gone.docx", storage_key="kb/x/never-written.docx",
        size_bytes=10, content_type="application/vnd.openxmlformats-officedocument"
                                    ".wordprocessingml.document",
    )
    db.add(doc)
    db.commit()

    res = client.get(f"/api/kb-documents/{doc.id}/text", headers=sys_headers)
    assert res.status_code == 200
    assert res.json()["available"] is False


def test_the_preview_needs_a_signed_in_user(client, world, files_dir):
    sys_headers = auth(client, "sys@x.org")
    doc_id = _upload(client, sys_headers, filename="p.docx",
                     content=_docx("anything")).json()["id"]
    assert client.get(f"/api/kb-documents/{doc_id}/text").status_code == 401


# --- the library is curated centrally ----------------------------------------


def test_a_department_user_cannot_browse_the_library(client, world, files_dir):
    """Staff reach the knowledge base through the assistant, not by browsing.
    404 rather than 403, so the refusal reveals nothing about what exists."""
    admin = auth(client, "a1@x.org")
    _upload(client, admin, title="Waste Guide")

    user_headers = auth(client, "u1@x.org")
    assert client.get("/api/kb-documents", headers=user_headers).status_code == 404
    assert (
        client.get("/api/kb-documents?search=waste", headers=user_headers).status_code
        == 404
    )


def test_a_department_user_can_still_open_a_cited_document(client, world, files_dir):
    """The assistant answers from the library for everyone, so the document a
    citation points at has to open — otherwise its sources are unverifiable."""
    admin = auth(client, "a1@x.org")
    doc_id = _upload(client, admin, title="Waste Guide").json()["id"]

    res = client.get(f"/api/kb-documents/{doc_id}", headers=auth(client, "u1@x.org"))
    assert res.status_code == 200
    assert res.json()["title"] == "Waste Guide"


def test_both_kinds_of_administrator_can_browse(client, world, files_dir):
    _upload(client, auth(client, "a1@x.org"), title="Waste Guide")
    for email in ("a1@x.org", "sys@x.org"):
        res = client.get("/api/kb-documents", headers=auth(client, email))
        assert res.status_code == 200, email
        assert [d["title"] for d in res.json()] == ["Waste Guide"]


# --- per-municipality libraries -------------------------------------------
#
# Each municipality has its own library beside the shared one. The whole point
# is that it is theirs: another municipality must not see it in a listing, must
# not be able to open it by guessing its id, and — the one that actually
# matters — must never have the assistant answer them out of it.


@pytest.fixture()
def two_cities(db, world):
    from app.core.security import hash_password
    from app.models import Municipality, User

    m2 = Municipality(name="City Two")
    pw = hash_password("kb-password-111")
    db.add_all([
        m2,
        User(email="a2@x.org", role="municipality_admin", municipality=m2,
             status="active", password_hash=pw, name="Other Admin"),
        User(email="u2@x.org", role="department_user", municipality=m2,
             status="active", password_hash=pw, name="Other Worker"),
    ])
    db.commit()
    return m2


def test_municipality_admin_uploads_land_in_their_own_library(client, world, files_dir):
    r = _upload(client, auth(client, "a1@x.org"), title="City One Procedure")
    assert r.status_code == 201, r.text
    assert r.json()["scope"] == "municipality"
    assert r.json()["municipality_name"] == "City One"


def test_system_admin_uploads_to_the_shared_library_by_default(client, world, files_dir):
    r = _upload(client, auth(client, "sys@x.org"), title="Programme Handbook")
    assert r.status_code == 201, r.text
    assert r.json()["scope"] == "global"
    assert r.json()["municipality_id"] is None


def test_system_admin_can_upload_into_a_named_municipality(client, world, files_dir):
    r = client.post(
        "/api/kb-documents",
        files={"file": ("guide.docx", _docx(), DOCX_MIME)},
        data={"scope": "municipality", "municipality_id": str(world["m1"].id)},
        headers=auth(client, "sys@x.org"),
    )
    assert r.status_code == 201, r.text
    assert r.json()["scope"] == "municipality"
    assert r.json()["municipality_name"] == "City One"


def test_a_municipality_library_is_invisible_to_another_municipality(
    client, world, two_cities, files_dir
):
    doc_id = _upload(client, auth(client, "a1@x.org"), title="City One Only").json()["id"]

    other = auth(client, "a2@x.org")
    assert client.get("/api/kb-documents", headers=other).json() == []
    # 404, not 403: the error must not confirm the document exists
    assert client.get(f"/api/kb-documents/{doc_id}", headers=other).status_code == 404
    assert client.get(f"/api/kb-documents/{doc_id}/text", headers=other).status_code == 404
    assert client.delete(f"/api/kb-documents/{doc_id}", headers=other).status_code == 404

    # and not to that municipality's staff either
    worker = auth(client, "u2@x.org")
    assert client.get(f"/api/kb-documents/{doc_id}", headers=worker).status_code == 404


def test_shared_library_stays_readable_by_everyone(client, world, two_cities, files_dir):
    doc_id = _upload(client, auth(client, "sys@x.org"), title="Shared").json()["id"]
    for email in ("a1@x.org", "u1@x.org", "a2@x.org", "u2@x.org"):
        r = client.get(f"/api/kb-documents/{doc_id}", headers=auth(client, email))
        assert r.status_code == 200, email


def test_municipality_admin_cannot_touch_the_shared_library(client, world, files_dir):
    doc_id = _upload(client, auth(client, "sys@x.org"), title="Shared").json()["id"]
    admin = auth(client, "a1@x.org")
    # readable — it is shared — but not theirs to change
    assert client.get(f"/api/kb-documents/{doc_id}", headers=admin).status_code == 200
    assert client.delete(f"/api/kb-documents/{doc_id}", headers=admin).status_code == 404
    assert client.post(f"/api/kb-documents/{doc_id}/retry", headers=admin).status_code == 404


def test_any_admin_of_the_municipality_manages_its_library(client, db, world, files_dir):
    """Not only whoever uploaded — a library must survive an administrator leaving."""
    from app.core.security import hash_password
    from app.models import User

    db.add(User(email="a1b@x.org", role="municipality_admin", municipality=world["m1"],
                status="active", password_hash=hash_password("kb-password-111"),
                name="Second Admin"))
    db.commit()

    doc_id = _upload(client, auth(client, "a1@x.org")).json()["id"]
    assert client.delete(
        f"/api/kb-documents/{doc_id}", headers=auth(client, "a1b@x.org")
    ).status_code == 200


def test_a_municipality_document_is_indexed_at_municipality_visibility(
    client, db, world, files_dir
):
    """The chunk's visibility is what retrieval filters on — global would leak it."""
    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    _upload(client, auth(client, "a1@x.org"))
    run_pending_jobs(db)

    chunk = db.query(Chunk).one()
    assert chunk.visibility == "municipality"
    assert chunk.municipality_id == world["m1"].id


def test_the_assistant_never_cites_another_municipalitys_library(
    client, db, world, two_cities, files_dir
):
    from app.rag.retrieval import retrieve
    from app.services.ingestion import run_pending_jobs

    _upload(
        client,
        auth(client, "a1@x.org"),
        content=_docx("Refuse collection in City One runs every Tuesday morning."),
    )
    run_pending_jobs(db)

    from app.models import User
    from app.rag.embeddings import get_embedding_provider

    question = "When is refuse collected?"
    vec = get_embedding_provider().embed([question])[0]
    theirs = db.query(User).filter_by(email="u1@x.org").one()
    outsider = db.query(User).filter_by(email="u2@x.org").one()

    def hits(user):
        return retrieve(db, query_embedding=vec, user=user, query_text=question)

    assert hits(theirs), "their own document must be reachable"
    assert hits(outsider) == []


def test_reembedding_keeps_a_document_in_its_own_library(client, db, world, files_dir):
    """Re-embedding must read each document's scope, not assume the shared one.

    Switching embedding model re-queues every document. If that queued a
    municipality's library as global, one maintenance command would publish
    every municipality's private material to all the others — silently, with
    nothing in the interface to show it had happened.
    """
    import sys
    from pathlib import Path

    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    _upload(client, auth(client, "a1@x.org"))
    _upload(client, auth(client, "sys@x.org"), title="Shared")
    run_pending_jobs(db)

    sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
    import reembed

    reembed.requeue(db)
    db.commit()
    run_pending_jobs(db)

    by_visibility = {c.visibility for c in db.query(Chunk).all()}
    assert by_visibility == {"municipality", "global"}
    muni_chunk = db.query(Chunk).filter_by(visibility="municipality").one()
    assert muni_chunk.municipality_id == world["m1"].id
