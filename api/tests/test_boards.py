import io

import pytest


@pytest.fixture()
def files_dir(tmp_path, monkeypatch):
    from app.core.config import get_settings

    monkeypatch.setattr(get_settings(), "files_dir", str(tmp_path))
    return tmp_path


@pytest.fixture()
def world(db):
    from app.core.security import hash_password
    from app.models import Category, Department, Municipality, User

    m1 = Municipality(name="City One")
    m2 = Municipality(name="City Two")
    d1 = Department(municipality=m1, name="Welfare")
    cat = Category(name_he="כלים", name_en="Tools")
    pw = hash_password("board-password-1")
    sysadmin = User(email="sys@x.org", role="system_admin", status="active",
                    password_hash=pw, name="Root")
    a1 = User(email="a1@x.org", role="municipality_admin", municipality=m1,
              status="active", password_hash=pw, name="Admin One")
    u1 = User(email="u1@x.org", role="department_user", municipality=m1,
              status="active", password_hash=pw, name="Worker One",
              departments=[d1])
    u2 = User(email="u2@x.org", role="department_user", municipality=m2,
              status="active", password_hash=pw, name="Worker Two")
    db.add_all([m1, m2, d1, cat, sysadmin, a1, u1, u2])
    db.commit()
    return {"m1": m1, "m2": m2, "d1": d1, "cat": cat, "u1": u1, "u2": u2, "a1": a1}


def auth(client, email):
    r = client.post(
        "/api/auth/login", json={"email": email, "password": "board-password-1"}
    )
    assert r.status_code == 200
    return {"Authorization": f"Bearer {r.json()['access_token']}"}


def publish_link(client, headers, world, *, title="GPT guide", destination="global"):
    return client.post(
        "/api/board-items",
        data={
            "title": title,
            "category_id": str(world["cat"].id),
            "destination": destination,
            "description": "How to write summaries",
            "link_url": "https://example.org/guide",
        },
        headers=headers,
    )


def test_publish_link_and_list(client, db, world):
    headers = auth(client, "u1@x.org")
    r = publish_link(client, headers, world)
    assert r.status_code == 201, r.text
    item = r.json()
    assert item["scope"] == "global" and item["link_url"].startswith("https://")
    assert item["author"]["municipality_name"] == "City One"

    page = client.get("/api/board-items?scope=global", headers=headers).json()
    assert [i["title"] for i in page["items"]] == ["GPT guide"]
    assert page["has_more"] is False

    # description was queued for indexing
    from app.models import IngestionJob

    assert db.query(IngestionJob).filter(IngestionJob.source_type == "board").count() == 1


def test_a_post_must_carry_something_and_links_must_be_https(client, world):
    headers = auth(client, "u1@x.org")
    r = client.post(
        "/api/board-items",
        data={
            "title": "No content",
            "category_id": str(world["cat"].id),
            "destination": "global",
        },
        headers=headers,
    )
    assert r.status_code == 422
    r = client.post(
        "/api/board-items",
        data={
            "title": "http link",
            "category_id": str(world["cat"].id),
            "destination": "global",
            "link_url": "http://insecure.example",
        },
        headers=headers,
    )
    assert r.status_code == 422


def test_municipality_board_scoping_404(client, world):
    headers1 = auth(client, "u1@x.org")
    r = publish_link(client, headers1, world, title="City One only",
                     destination="municipality")
    assert r.status_code == 201
    item_id = r.json()["id"]
    assert r.json()["municipality_name" if False else "scope"] == "municipality"

    # member of another municipality: list empty + direct fetch 404
    headers2 = auth(client, "u2@x.org")
    page = client.get("/api/board-items?scope=municipality", headers=headers2).json()
    assert page["items"] == []
    assert (
        client.get(f"/api/board-items/{item_id}", headers=headers2).status_code == 404
    )
    # sysadmin can fetch
    assert (
        client.get(f"/api/board-items/{item_id}", headers=auth(client, "sys@x.org"))
        .status_code
        == 200
    )


def test_like_toggle_and_comments(client, world):
    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]

    r = client.post(f"/api/board-items/{item_id}/like", headers=headers)
    assert r.json() == {"liked": True, "like_count": 1}
    r = client.post(f"/api/board-items/{item_id}/like", headers=headers)
    assert r.json() == {"liked": False, "like_count": 0}

    r = client.post(
        f"/api/board-items/{item_id}/comments", json={"body": "great tool"},
        headers=headers,
    )
    assert r.status_code == 201
    comment_id = r.json()["id"]
    detail = client.get(f"/api/board-items/{item_id}", headers=headers).json()
    assert [c["body"] for c in detail["comments"]] == ["great tool"]

    # over-long comment rejected
    r = client.post(
        f"/api/board-items/{item_id}/comments", json={"body": "x" * 1001},
        headers=headers,
    )
    assert r.status_code == 422

    assert (
        client.delete(
            f"/api/board-items/{item_id}/comments/{comment_id}", headers=headers
        ).status_code
        == 200
    )


def test_edit_author_only(client, world):
    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    patch = {
        "title": "Renamed", "description": "new desc",
        "category_id": str(world["cat"].id),
    }
    other = auth(client, "u2@x.org")
    assert (
        client.patch(f"/api/board-items/{item_id}", json=patch, headers=other)
        .status_code
        == 404
    )
    r = client.patch(f"/api/board-items/{item_id}", json=patch, headers=headers)
    assert r.status_code == 200 and r.json()["title"] == "Renamed"


def test_delete_rules_and_moderation_audit(client, db, world):
    author = auth(client, "u1@x.org")
    item_id = publish_link(client, author, world).json()["id"]

    # unrelated user (other municipality) cannot delete
    assert (
        client.delete(f"/api/board-items/{item_id}", headers=auth(client, "u2@x.org"))
        .status_code
        == 404
    )
    # author's municipality admin CAN delete the global item (moderation) + audit row
    r = client.delete(f"/api/board-items/{item_id}", headers=auth(client, "a1@x.org"))
    assert r.status_code == 200

    from app.models import AuditLog

    assert (
        db.query(AuditLog)
        .filter(AuditLog.action == "board_item.moderate_delete")
        .count()
        == 1
    )


def test_file_item_upload_and_delete_cascades(client, db, world, files_dir):
    import docx

    from app.models import Chunk, IngestionJob
    from app.services.ingestion import run_pending_jobs

    d = docx.Document()
    d.add_paragraph("Annual budget template for municipalities")
    buf = io.BytesIO()
    d.save(buf)

    headers = auth(client, "u1@x.org")
    r = client.post(
        "/api/board-items",
        data={
            "title": "Budget template",
            "category_id": str(world["cat"].id),
            "destination": "global",
        },
        files={
            "file": (
                "budget.docx", buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=headers,
    )
    assert r.status_code == 201, r.text
    item_id = r.json()["id"]
    run_pending_jobs(db)
    chunks = db.query(Chunk).filter(Chunk.source_type == "board").all()
    assert chunks and any("budget" in c.content.lower() for c in chunks)

    detail = client.get(f"/api/board-items/{item_id}", headers=headers).json()
    assert detail["download_url"].startswith("/api/files/board/")

    assert client.delete(f"/api/board-items/{item_id}", headers=headers).status_code == 200
    assert db.query(Chunk).filter(Chunk.source_type == "board").count() == 0
    assert db.query(IngestionJob).filter(IngestionJob.source_type == "board").count() == 0


def test_search_and_category_filter(client, db, world):
    from app.models import Category

    headers = auth(client, "u1@x.org")
    other_cat = Category(name_he="טפסים", name_en="Forms")
    db.add(other_cat)
    db.commit()
    publish_link(client, headers, world, title="Hebrew guide שיטות עבודה")
    publish_link(client, headers, world, title="Budget forms")

    page = client.get("/api/board-items?search=שיטות", headers=headers).json()
    assert [i["title"] for i in page["items"]] == ["Hebrew guide שיטות עבודה"]
    page = client.get(
        f"/api/board-items?category_id={other_cat.id}", headers=headers
    ).json()
    assert page["items"] == []



# --- shared prompts and agents -----------------------------------------------


def publish_prompt(client, headers, world, **extra):
    data = {
        "title": "פרומפט לניסוח מכתב לתושב",
        "category_id": str(world["cat"].id),
        "destination": "global",
        "prompt_text": "אתה עוזר בעירייה. נסח תשובה מנומסת בעברית לפניית תושב.",
    }
    data.update(extra)
    return client.post("/api/board-items", data=data, headers=headers)


def test_publish_a_prompt_with_no_file_and_no_link(client, world):
    """A shared prompt is content in its own right. The previous rule demanded a
    file or a link, which would have rejected the most useful kind of post."""
    res = publish_prompt(client, auth(client, "u1@x.org"), world)
    assert res.status_code == 201, res.text
    body = res.json()
    assert "אתה עוזר בעירייה" in body["prompt_text"]
    assert body["link_url"] is None and body["filename"] is None


def test_a_prompt_may_travel_with_a_link_to_where_the_agent_lives(client, world):
    res = publish_prompt(
        client, auth(client, "u1@x.org"), world,
        title="Resident enquiry agent",
        link_url="https://example.org/agents/resident-enquiry",
    )
    assert res.status_code == 201, res.text
    body = res.json()
    assert body["prompt_text"] and body["link_url"]


def test_a_shared_prompt_is_findable_by_the_assistant(client, db, world):
    """Storing the prompt as text is what makes it searchable — otherwise a
    colleague can only find it by scrolling the board."""
    import uuid

    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    res = publish_prompt(
        client, auth(client, "u1@x.org"), world,
        title="Tender summary prompt",
        prompt_text="Summarise a tender document under regulation 17.3 in Hebrew",
    )
    assert res.status_code == 201
    item_id = uuid.UUID(res.json()["id"])

    run_pending_jobs(db)
    chunks = db.query(Chunk).filter(Chunk.source_id == item_id).all()
    assert chunks, "the prompt was never indexed"
    assert any("regulation 17.3" in c.content for c in chunks)


def test_any_file_type_can_be_attached_to_a_post(client, db, world, files_dir):
    """Unrestricted uploads are a product decision; an unreadable type must not
    fail the post."""
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    res = client.post(
        "/api/board-items",
        data={
            "title": "Team screenshot",
            "category_id": str(world["cat"].id),
            "destination": "global",
        },
        files={"file": ("diagram.zip", b"PK\x03\x04nonsense", "application/zip")},
        headers=headers,
    )
    assert res.status_code == 201, res.text
    assert res.json()["filename"] == "diagram.zip"

    run_pending_jobs(db)  # must not raise on a type it cannot read


def test_board_search_covers_title_description_and_prompt(client, db, world):
    """Search is how people find a post they half-remember. A shared prompt's
    body is the substance of that post, so it has to be searchable too."""
    headers = auth(client, "u1@x.org")
    cid = str(world["cat"].id)

    client.post("/api/board-items", headers=headers, data={
        "title": "Tender checklist", "category_id": cid, "destination": "global",
        "link_url": "https://example.org/x"})
    client.post("/api/board-items", headers=headers, data={
        "title": "Meeting notes", "category_id": cid, "destination": "global",
        "description": "Covers the recycling centre opening hours",
        "link_url": "https://example.org/y"})
    client.post("/api/board-items", headers=headers, data={
        "title": "נוהל פנימי", "category_id": cid, "destination": "global",
        "description": "סיכום ישיבת ועדת התקציב", "link_url": "https://example.org/z"})
    client.post("/api/board-items", headers=headers, data={
        "title": "Helper", "category_id": cid, "destination": "global",
        "prompt_text": "Summarise a tender document under regulation 17.3"})

    def titles(term):
        r = client.get(f"/api/board-items?scope=global&search={term}", headers=headers)
        assert r.status_code == 200, r.text
        return [i["title"] for i in r.json()["items"]]

    assert "Tender checklist" in titles("tender")          # title
    assert "Meeting notes" in titles("recycling")           # description
    assert "נוהל פנימי" in titles("התקציב")                  # Hebrew
    assert "Helper" in titles("regulation"), "prompt body is not searchable"


# --- comment replies ----------------------------------------------------------


def test_reply_attaches_to_its_parent(client, world):
    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    parent = client.post(f"/api/board-items/{item_id}/comments",
                         json={"body": "does this cover tenders?"}, headers=headers).json()

    reply = client.post(f"/api/board-items/{item_id}/comments",
                        json={"body": "yes, section 4", "parent_id": parent["id"]},
                        headers=headers)
    assert reply.status_code == 201, reply.text
    assert reply.json()["parent_id"] == parent["id"]

    detail = client.get(f"/api/board-items/{item_id}", headers=headers).json()
    by_id = {c["id"]: c for c in detail["comments"]}
    assert by_id[parent["id"]]["parent_id"] is None
    assert by_id[reply.json()["id"]]["parent_id"] == parent["id"]


def test_replying_to_a_reply_stays_one_level_deep(client, world):
    """Deeper threads are hard to read and rarely say more than a flat reply."""
    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    parent = client.post(f"/api/board-items/{item_id}/comments",
                         json={"body": "top"}, headers=headers).json()
    reply = client.post(f"/api/board-items/{item_id}/comments",
                        json={"body": "middle", "parent_id": parent["id"]},
                        headers=headers).json()

    deeper = client.post(f"/api/board-items/{item_id}/comments",
                         json={"body": "deeper", "parent_id": reply["id"]},
                         headers=headers)
    assert deeper.status_code == 201
    assert deeper.json()["parent_id"] == parent["id"], "nesting went past one level"


def test_reply_cannot_attach_to_another_posts_comment(client, world):
    """Otherwise a reply lands somewhere its author never looks."""
    headers = auth(client, "u1@x.org")
    first = publish_link(client, headers, world, title="First").json()["id"]
    second = publish_link(client, headers, world, title="Second").json()["id"]
    elsewhere = client.post(f"/api/board-items/{first}/comments",
                            json={"body": "on the first post"}, headers=headers).json()

    res = client.post(f"/api/board-items/{second}/comments",
                      json={"body": "misdirected", "parent_id": elsewhere["id"]},
                      headers=headers)
    assert res.status_code == 422
    assert res.json()["detail"] == "invalid_parent"


def test_deleting_a_comment_removes_its_replies(client, world):
    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    parent = client.post(f"/api/board-items/{item_id}/comments",
                         json={"body": "top"}, headers=headers).json()
    client.post(f"/api/board-items/{item_id}/comments",
                json={"body": "a reply", "parent_id": parent["id"]}, headers=headers)

    assert client.delete(f"/api/board-items/{item_id}/comments/{parent['id']}",
                         headers=headers).status_code == 200
    detail = client.get(f"/api/board-items/{item_id}", headers=headers).json()
    assert detail["comments"] == [], "an orphaned reply survived its parent"


# --- comment reactions --------------------------------------------------------


def _comment_on(client, headers, world, body="worth reading"):
    item_id = publish_link(client, headers, world).json()["id"]
    cid = client.post(f"/api/board-items/{item_id}/comments",
                      json={"body": body}, headers=headers).json()["id"]
    return item_id, cid


def test_reaction_toggles_on_and_off(client, world):
    headers = auth(client, "u1@x.org")
    item_id, cid = _comment_on(client, headers, world)
    url = f"/api/board-items/{item_id}/comments/{cid}/reactions"

    on = client.post(url, json={"emoji": "👍"}, headers=headers)
    assert on.status_code == 200, on.text
    assert on.json() == {"emoji": "👍", "count": 1, "mine": True}

    off = client.post(url, json={"emoji": "👍"}, headers=headers)
    assert off.json() == {"emoji": "👍", "count": 0, "mine": False}


def test_the_same_person_cannot_double_count_one_emoji(client, world):
    """The key is (comment, person, emoji), so a double click toggles rather
    than accumulating."""
    headers = auth(client, "u1@x.org")
    item_id, cid = _comment_on(client, headers, world)
    url = f"/api/board-items/{item_id}/comments/{cid}/reactions"

    client.post(url, json={"emoji": "🎉"}, headers=headers)
    client.post(url, json={"emoji": "🎉"}, headers=headers)
    third = client.post(url, json={"emoji": "🎉"}, headers=headers)
    assert third.json()["count"] == 1


def test_several_people_and_several_emoji_are_counted_separately(client, world):
    one = auth(client, "u1@x.org")
    two = auth(client, "sys@x.org")
    item_id, cid = _comment_on(client, one, world)
    url = f"/api/board-items/{item_id}/comments/{cid}/reactions"

    client.post(url, json={"emoji": "👍"}, headers=one)
    client.post(url, json={"emoji": "👍"}, headers=two)
    client.post(url, json={"emoji": "❤️"}, headers=one)

    detail = client.get(f"/api/board-items/{item_id}", headers=one).json()
    reactions = {r["emoji"]: r for r in detail["comments"][0]["reactions"]}
    assert reactions["👍"]["count"] == 2 and reactions["👍"]["mine"] is True
    assert reactions["❤️"]["count"] == 1

    # and "mine" is per person, not global
    detail_two = client.get(f"/api/board-items/{item_id}", headers=two).json()
    by_emoji = {r["emoji"]: r for r in detail_two["comments"][0]["reactions"]}
    assert by_emoji["❤️"]["mine"] is False


def test_only_the_supported_emoji_are_accepted(client, world):
    """The value is rendered directly, so the column is not free text."""
    headers = auth(client, "u1@x.org")
    item_id, cid = _comment_on(client, headers, world)
    url = f"/api/board-items/{item_id}/comments/{cid}/reactions"

    for bad in ["💩", "<script>", "👍👍", ""]:
        assert client.post(url, json={"emoji": bad}, headers=headers).status_code == 422


def test_reaction_on_another_posts_comment_is_refused(client, world):
    headers = auth(client, "u1@x.org")
    _, cid = _comment_on(client, headers, world)
    other_item = publish_link(client, headers, world, title="Other").json()["id"]

    res = client.post(f"/api/board-items/{other_item}/comments/{cid}/reactions",
                      json={"emoji": "👍"}, headers=headers)
    assert res.status_code == 404


def test_deleting_a_comment_removes_its_reactions(client, db, world):
    from app.models import BoardCommentReaction

    headers = auth(client, "u1@x.org")
    item_id, cid = _comment_on(client, headers, world)
    client.post(f"/api/board-items/{item_id}/comments/{cid}/reactions",
                json={"emoji": "🙏"}, headers=headers)
    assert db.query(BoardCommentReaction).count() == 1

    client.delete(f"/api/board-items/{item_id}/comments/{cid}", headers=headers)
    assert db.query(BoardCommentReaction).count() == 0


def test_the_assistant_can_use_board_posts_and_respects_their_scope(client, db, world):
    """A board post is knowledge too. It should be quotable by the assistant —
    and a municipality post must stay inside its municipality when it is."""
    from app.models import User
    from app.rag.embeddings import get_embedding_provider
    from app.rag.retrieval import retrieve
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    client.post("/api/board-items", headers=headers, data={
        "title": "Recycling centre hours",
        "category_id": str(world["cat"].id),
        "destination": "global",
        "description": "The recycling centre opens Sunday to Thursday, 08:00 to 16:00",
        "link_url": "https://example.org/recycling"})
    client.post("/api/board-items", headers=headers, data={
        "title": "City One procurement note",
        "category_id": str(world["cat"].id),
        "destination": "municipality",
        "description": "City One uses form 7Q for internal procurement requests",
        "link_url": "https://example.org/procurement"})
    run_pending_jobs(db)

    def ask(email, question):
        user = db.query(User).filter_by(email=email).one()
        [vec] = get_embedding_provider().embed([question])
        hits = retrieve(db, query_embedding=vec, user=user, query_text=question)
        return " ".join(h.content for h in hits)

    # global post reaches everyone, including another municipality
    assert "recycling centre opens" in ask("u2@x.org", "When does the recycling centre open?")

    # municipality post reaches its own members
    assert "form 7Q" in ask("u1@x.org", "Which form is used for procurement requests?")
    # but not another municipality
    assert "form 7Q" not in ask("u2@x.org", "Which form is used for procurement requests?")


def test_a_shared_prompt_post_is_quotable_by_the_assistant(client, db, world):
    from app.models import User
    from app.rag.embeddings import get_embedding_provider
    from app.rag.retrieval import retrieve
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    client.post("/api/board-items", headers=headers, data={
        "title": "Tender summary helper",
        "category_id": str(world["cat"].id),
        "destination": "global",
        "prompt_text": "Summarise a tender document under regulation 17.3 in Hebrew"})
    run_pending_jobs(db)

    user = db.query(User).filter_by(email="u2@x.org").one()
    q = "regulation 17.3 tender"
    [vec] = get_embedding_provider().embed([q])
    hits = retrieve(db, query_embedding=vec, user=user, query_text=q)
    assert any("regulation 17.3" in h.content for h in hits)
    assert any(h.source_type == "board" for h in hits)


def test_a_file_attached_to_a_post_is_readable_by_the_assistant(
    client, db, world, files_dir
):
    """A post's attachment is knowledge too — the words inside the file, not
    just the post's own description."""
    import docx

    from app.models import User
    from app.rag.embeddings import get_embedding_provider
    from app.rag.retrieval import retrieve
    from app.services.ingestion import run_pending_jobs

    d = docx.Document()
    d.add_paragraph("Waste collection operates on Mondays and Thursdays in the north")
    buf = io.BytesIO()
    d.save(buf)

    headers = auth(client, "u1@x.org")
    res = client.post(
        "/api/board-items",
        data={
            "title": "Collection schedule",
            "category_id": str(world["cat"].id),
            "destination": "global",
            "description": "Attached is this year's schedule",
        },
        files={"file": ("schedule.docx", buf.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=headers,
    )
    assert res.status_code == 201, res.text
    run_pending_jobs(db)

    user = db.query(User).filter_by(email="u2@x.org").one()
    q = "Which days does waste collection operate in the north?"
    [vec] = get_embedding_provider().embed([q])
    hits = retrieve(db, query_embedding=vec, user=user, query_text=q)
    assert any("Mondays and Thursdays" in h.content for h in hits), (
        "the attachment's contents were not indexed"
    )


# --- comments as knowledge ----------------------------------------------------


def test_a_comment_is_answerable_and_cites_its_post(client, db, world):
    """Corrections live in the replies — 'that form was replaced' — and were
    invisible to the assistant until comments were indexed."""
    from app.models import User
    from app.rag.embeddings import get_embedding_provider
    from app.rag.retrieval import retrieve
    from app.services.citations import build_citations
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world, title="Procurement forms").json()["id"]
    client.post(f"/api/board-items/{item_id}/comments",
                json={"body": "Note: form 7Q was replaced by form 8R in March"},
                headers=headers)
    run_pending_jobs(db)

    user = db.query(User).filter_by(email="u2@x.org").one()
    q = "form 8R replaced"
    [vec] = get_embedding_provider().embed([q])
    hits = retrieve(db, query_embedding=vec, user=user, query_text=q)
    assert any("form 8R" in h.content for h in hits), "the comment was not indexed"

    # the citation points at the post, since a comment has no page of its own
    cites = build_citations(db, [h for h in hits if h.source_type == "comment"])
    assert cites and cites[0]["href"] == f"/board/{item_id}"


def test_a_comment_never_reaches_further_than_its_post(client, db, world):
    from app.models import User
    from app.rag.embeddings import get_embedding_provider
    from app.rag.retrieval import retrieve
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world, title="Internal note",
                           destination="municipality").json()["id"]
    client.post(f"/api/board-items/{item_id}/comments",
                json={"body": "City One uses reference 9K internally"}, headers=headers)
    run_pending_jobs(db)

    def sees(email):
        user = db.query(User).filter_by(email=email).one()
        [vec] = get_embedding_provider().embed(["reference 9K"])
        hits = retrieve(db, query_embedding=vec, user=user, query_text="reference 9K")
        return any("9K" in h.content for h in hits)

    assert sees("u1@x.org")           # same municipality
    assert not sees("u2@x.org")       # another municipality


def test_deleting_a_comment_removes_it_from_the_assistant(client, db, world):
    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    cid = client.post(f"/api/board-items/{item_id}/comments",
                      json={"body": "a passing remark"}, headers=headers).json()["id"]
    run_pending_jobs(db)
    assert db.query(Chunk).filter(Chunk.source_type == "comment").count() == 1

    client.delete(f"/api/board-items/{item_id}/comments/{cid}", headers=headers)
    assert db.query(Chunk).filter(Chunk.source_type == "comment").count() == 0


def test_deleting_a_post_removes_its_comments_from_the_assistant(client, db, world):
    """Comment chunks are keyed by comment id, so they would otherwise survive
    the post — answerable content for something that no longer exists."""
    from app.models import Chunk
    from app.services.ingestion import run_pending_jobs

    headers = auth(client, "u1@x.org")
    item_id = publish_link(client, headers, world).json()["id"]
    client.post(f"/api/board-items/{item_id}/comments",
                json={"body": "first"}, headers=headers)
    client.post(f"/api/board-items/{item_id}/comments",
                json={"body": "second"}, headers=headers)
    run_pending_jobs(db)
    assert db.query(Chunk).filter(Chunk.source_type == "comment").count() == 2

    client.delete(f"/api/board-items/{item_id}", headers=headers)
    assert db.query(Chunk).filter(Chunk.source_type == "comment").count() == 0


# --- images -----------------------------------------------------------------

# The smallest valid PNG: one transparent pixel.
ONE_PIXEL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06"
    b"\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05"
    b"\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def publish_image(client, headers, world, *, filename="poster.png",
                  content=ONE_PIXEL_PNG, content_type="image/png"):
    return client.post(
        "/api/board-items",
        data={
            "title": "Summer programme poster",
            "category_id": str(world["cat"].id),
            "destination": "global",
        },
        files={"file": (filename, content, content_type)},
        headers=headers,
    )


def test_an_image_comes_back_ready_to_show(client, files_dir, world):
    """Otherwise a poster is a download button with a filename on it."""
    headers = auth(client, "u1@x.org")
    r = publish_image(client, headers, world)
    assert r.status_code == 201, r.text
    assert r.json()["image_url"]

    # and in the list, so a card can show it without opening the post
    page = client.get("/api/board-items?scope=global", headers=headers).json()
    listed = [i for i in page["items"] if i["id"] == r.json()["id"]]
    assert listed and listed[0]["image_url"]


def test_a_document_is_not_offered_as_an_image(client, files_dir, world):
    """A PDF may be shown in a viewer; it has no business stretched across a
    card."""
    headers = auth(client, "u1@x.org")
    r = client.post(
        "/api/board-items",
        data={
            "title": "Procurement rules",
            "category_id": str(world["cat"].id),
            "destination": "global",
        },
        files={"file": ("rules.pdf", b"%PDF-1.4 rules", "application/pdf")},
        headers=headers,
    )
    assert r.status_code == 201, r.text
    assert r.json()["image_url"] is None


def test_a_post_with_no_file_has_no_image(client, world):
    r = publish_link(client, auth(client, "u1@x.org"), world)
    assert r.json()["image_url"] is None


def test_an_svg_is_never_shown_as_an_image(client, files_dir, world):
    """It can carry script. It uploads like anything else and downloads like
    anything else, but it is never dropped into the page."""
    headers = auth(client, "u1@x.org")
    r = publish_image(
        client, headers, world,
        filename="logo.svg",
        content=b"<svg xmlns='http://www.w3.org/2000/svg'><script>alert(1)</script></svg>",
        content_type="image/svg+xml",
    )
    assert r.status_code == 201, r.text
    assert r.json()["image_url"] is None


# --- a system admin reading a municipality board ----------------------------


def test_a_system_admin_can_read_one_municipality_board(client, world):
    """They could already open any single post on one; the list is the part
    that was missing."""
    r = publish_link(client, auth(client, "u1@x.org"), world,
                     title="City One only", destination="municipality")
    assert r.status_code == 201

    sys_headers = auth(client, "sys@x.org")
    page = client.get(
        f"/api/board-items?scope=municipality&municipality_id={world['m1'].id}",
        headers=sys_headers,
    ).json()
    assert [i["title"] for i in page["items"]] == ["City One only"]

    # and the other municipality's board is genuinely a different board
    other = client.get(
        f"/api/board-items?scope=municipality&municipality_id={world['m2'].id}",
        headers=sys_headers,
    ).json()
    assert other["items"] == []


def test_naming_no_municipality_gives_a_system_admin_all_of_them(client, world):
    """"Is anyone asking about this?" is a question about every board at once,
    not about one at a time."""
    sys_headers = auth(client, "sys@x.org")
    publish_link(client, auth(client, "u1@x.org"), world,
                 title="City One only", destination="municipality")
    publish_link(client, auth(client, "u2@x.org"), world,
                 title="City Two only", destination="municipality")

    page = client.get("/api/board-items?scope=municipality", headers=sys_headers).json()
    assert sorted(i["title"] for i in page["items"]) == ["City One only", "City Two only"]
    # and still only the municipality boards — the shared one is its own page
    publish_link(client, auth(client, "u1@x.org"), world, title="Everyone")
    page = client.get("/api/board-items?scope=municipality", headers=sys_headers).json()
    assert "Everyone" not in [i["title"] for i in page["items"]]


def test_somebody_with_no_municipality_still_has_no_board(client, db, world):
    """The all-boards view is the system admin's, not a consolation prize for
    an account that was never given a municipality."""
    from app.core.security import hash_password
    from app.models import User

    db.add(User(email="nowhere@x.org", role="department_user", status="active",
                password_hash=hash_password("board-password-1"), name="Nowhere"))
    db.commit()

    r = client.get("/api/board-items?scope=municipality",
                   headers=auth(client, "nowhere@x.org"))
    assert r.status_code == 404


def test_naming_somebody_elses_municipality_changes_nothing(client, world):
    """Same 404 as a municipality that does not exist, so no one learns which
    ids are real."""
    r = client.get(
        f"/api/board-items?scope=municipality&municipality_id={world['m2'].id}",
        headers=auth(client, "u1@x.org"),
    )
    assert r.status_code == 404


def test_naming_your_own_municipality_is_allowed(client, world):
    """The web sends it for everyone rather than branching on the role."""
    r = client.get(
        f"/api/board-items?scope=municipality&municipality_id={world['m1'].id}",
        headers=auth(client, "u1@x.org"),
    )
    assert r.status_code == 200


def test_a_made_up_municipality_id_is_a_404_not_a_crash(client, world):
    r = client.get(
        "/api/board-items?scope=municipality&municipality_id=not-a-uuid",
        headers=auth(client, "sys@x.org"),
    )
    assert r.status_code == 404
