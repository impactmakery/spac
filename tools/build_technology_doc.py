"""Build *Tomorrow Agent Hub — Technology* as a .docx.

The same parts list as TECHNOLOGY.md, in a form that can be handed to someone
who does not read the repository. Generated rather than hand-edited: editing
the .docx directly means the next build silently discards the change.

    python tools/build_technology_doc.py                    # to ~/Downloads
    python tools/build_technology_doc.py --out ./tech.docx

Requires python-docx, which the API already depends on:

    api/.venv/Scripts/python tools/build_technology_doc.py   # Windows
    api/.venv/bin/python tools/build_technology_doc.py       # macOS / Linux
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DEFAULT_OUT = Path.home() / "Downloads" / "Tomorrow Agent Hub - Technology.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)


def main() -> None:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(description="Build the technology document.")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def h(text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level)
        for run in p.runs:
            run.font.color.rgb = ACCENT

    def para(text: str = "", *, muted: bool = False, italic: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = italic
        if muted:
            run.font.color.rgb = MUTED
            run.font.size = Pt(9.5)

    def bullets(items: list[str]) -> None:
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            if " — " in item:
                head, _, tail = item.partition(" — ")
                p.add_run(head).bold = True
                p.add_run(" — " + tail)
            else:
                p.add_run(item)

    def table(headers: list[str], rows: list[list[str]]) -> None:
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for cell, text in zip(t.rows[0].cells, headers, strict=True):
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for cell, text in zip(cells, row, strict=True):
                cell.text = ""
                cell.paragraphs[0].add_run(text)
        doc.add_paragraph()

    # ------------------------------------------------------------------ title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tomorrow Agent Hub")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technology")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    para()
    para(
        "Everything the platform is built from, what each piece is for, and why it "
        "was chosen over the obvious alternative.",
        muted=True,
    )
    para(
        "Generated from the repository. Versions are the minimum declared in "
        "api/requirements.txt and web/package.json, which are the source of truth "
        "if the two ever disagree.",
        muted=True,
    )

    props = doc.core_properties
    props.title = "Tomorrow Agent Hub - Technology"
    props.comments = (
        "Parts list for Tomorrow Agent Hub: backend, frontend, document reading, "
        "OCR, infrastructure, AI providers, search and tooling."
    )

    # ---------------------------------------------------------------- backend
    h("1. Backend — Python 3.12", 1)
    table(
        ["Package", "Version", "What it does"],
        [
            ["FastAPI", "≥0.115", "HTTP framework. Every permission check lives here."],
            ["Uvicorn", "≥0.30", "ASGI server."],
            ["SQLAlchemy", "≥2.0", "ORM and query builder. Retrieval drops to raw SQL where the permission filter has to sit inside the query."],
            ["Alembic", "≥1.13", "Migrations, applied automatically when a container starts. No manual schema changes."],
            ["psycopg", "≥3.2", "PostgreSQL driver."],
            ["pgvector", "≥0.3", "Vector column type and distance operators."],
            ["pydantic-settings", "≥2.4", "Configuration from the environment."],
            ["PyJWT", "≥2.9", "Session tokens."],
            ["bcrypt", "≥4.2", "Password hashing."],
            ["tiktoken", "≥0.7", "Token counting, so passages are 800 tokens rather than 800 characters."],
        ],
    )

    # --------------------------------------------------------------- frontend
    h("2. Frontend — Next.js 16, React 19, TypeScript 5", 1)
    table(
        ["Package", "Version", "What it does"],
        [
            ["Next.js", "16.3", "Pages fetch on the server, so the browser never holds an API token."],
            ["React", "19.2", "—"],
            ["next-intl", "≥4.13", "Hebrew and English. Every user-facing string lives in a translation file."],
            ["NextAuth (Auth.js)", "v5", "Session cookie. Credentials are verified by the API, never in the frontend."],
            ["Tailwind CSS", "v4", "Right-to-left comes from logical properties, so both languages share one stylesheet."],
            ["lucide-react", "≥1.28", "Icons."],
            ["Vitest", "≥4.1", "Unit tests."],
        ],
    )

    # ------------------------------------------------------- reading documents
    h("3. Reading documents", 1)
    para("One library per format, rather than one heavyweight document toolkit.")
    table(
        ["Package", "Format"],
        [
            ["pypdf ≥4.0", "PDF text layer"],
            ["python-docx ≥1.1", "Word"],
            ["python-pptx ≥0.6", "PowerPoint"],
            ["openpyxl ≥3.1", "Excel"],
            ["Pillow ≥10.0", "Images"],
        ],
    )

    h("OCR — reading words out of pictures", 2)
    para(
        "Used only when the libraries above return implausibly little text: a scanned "
        "PDF with no text layer, or an Office file whose content is images."
    )
    table(
        ["Package", "Role"],
        [
            ["Tesseract OCR", "The engine. Installed in our own container with Hebrew and English language packs."],
            ["pytesseract ≥0.3.13", "Python binding to it."],
            ["pypdfium2 ≥4.30", "Turns PDF pages into images so Tesseract has something to read."],
        ],
    )
    bullets(
        [
            "Costs no fees — Tesseract runs on our own server, so OCR costs processing time rather than money. It is why scanned documents are the slow part of any bulk upload.",
            "Hebrew quality — good on clean printed text, weak on handwriting, poor scans and stylised design work. If quality matters more than cost later, Google Document AI or Azure Document Intelligence are markedly better on Hebrew, priced per page.",
            "Licensing — pypdfium2 rather than the more common PyMuPDF, which is AGPL and would place a copyleft obligation on a commercial product.",
        ]
    )

    # --------------------------------------------------------- infrastructure
    doc.add_page_break()
    h("4. Infrastructure", 1)
    table(
        ["Service", "Role"],
        [
            ["PostgreSQL 16 + pgvector", "Application data and document embeddings in one database."],
            ["Cloudflare R2", "Uploaded files. Uploads go through the server; downloads use short-lived signed links."],
            ["Railway", "API, ingestion worker and scheduled jobs. Amsterdam region, chosen for latency to Israel."],
            ["Vercel", "The web application."],
            ["Resend", "Invitations, password resets and the weekly digest."],
            ["Sentry", "Error monitoring. Installed in the code but not yet switched on — no account is configured, so nothing is currently reported."],
        ],
    )

    # -------------------------------------------------------------------- AI
    h("5. Artificial intelligence", 1)
    para("Three separate jobs. Only one of them costs money.")
    table(
        ["Job", "Provider", "Model", "Cost"],
        [
            ["Answering questions", "OpenRouter", "google/gemma-4-26b-a4b-it (free tier), with two further free models as fallback", "None"],
            ["Understanding documents for search", "OpenAI", "text-embedding-3-large", "Paid, per document"],
            ["Building the knowledge graph", "None", "Pattern matching, no model is called", "None"],
        ],
    )
    bullets(
        [
            "Embedding is charged when a document is uploaded, not when someone asks a question — the whole 344-document municipal load cost roughly $0.20.",
            "Asking questions is currently free, which is also why there is a fallback chain: when a free model is rate-limited the request falls through to the next, and the reader sees a slower answer rather than an error.",
            "Answer quality is capped by those free models — finding the right passages and writing the answer are separate jobs. The search side is thoroughly tested; the writing side runs on a free model. Moving to a paid model is a one-line configuration change.",
            "The knowledge graph uses pattern matching rather than a model, switched during the first bulk load because building it with a model meant roughly ten thousand calls against a daily free quota. Search, ranking and citations are unaffected. Pointing it back at a funded model and re-indexing restores the fuller version.",
            "Changing the embedding model requires re-embedding everything — vectors from different models are not comparable, and mixing them quietly degrades search.",
        ]
    )

    # ---------------------------------------------------------------- search
    h("6. Search", 1)
    para(
        "No external search service. Three methods over the same database, combined "
        "and then re-ranked:"
    )
    bullets(
        [
            "Meaning — similarity between the question and each passage, using pgvector.",
            "Words — PostgreSQL full-text search, for the times someone quotes an exact phrase or a form number.",
            "Connections — traversal of the entities and relationships found when each document was indexed.",
        ]
    )
    para(
        "The permission filter sits inside every one of those queries rather than "
        "being applied to their results. That is the single most important design "
        "decision in the platform: a person can never receive an answer drawn from "
        "material they are not allowed to see."
    )

    # ------------------------------------------------------------ development
    h("7. Development and testing", 1)
    table(
        ["Tool", "Role"],
        [
            ["ruff", "Linting and import ordering."],
            ["mypy", "Type checking."],
            ["pytest", "Backend tests. The access-control suite must pass before anything merges."],
            ["pgserver", "Embedded PostgreSQL, so tests run without Docker."],
            ["GitHub Actions", "Runs type checking, linting and both test suites on every change."],
            ["Docker Compose", "Local database for development."],
        ],
    )
    para(
        "The test suite runs on deterministic offline stand-ins for the AI providers, "
        "so it is fast, free, and safe to run on a machine holding production keys.",
        muted=True,
    )

    # -------------------------------------------------- deliberately not used
    h("8. Deliberately not used", 1)
    bullets(
        [
            "No vector database — one PostgreSQL holds both application data and embeddings, so a permission check and a similarity search are the same query in the same transaction. A separate vector store would mean enforcing permissions twice, in two places, with two chances to disagree.",
            "No message broker — the job queue is a database table. It removes an entire service from the deployment and is more than enough at this scale.",
            "No LangChain, despite the original plan naming it — the pipeline is extract, split, embed, query, prompt, and each step is a few lines of explicit code. The abstraction would have obscured the permission filter, which is the part that most needs to be readable. Worth revisiting if multi-step agents are added.",
        ]
    )

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"wrote: {args.out}")


if __name__ == "__main__":
    main()
