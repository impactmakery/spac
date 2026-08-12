import argparse
import os
import sys
from pathlib import Path

# the sections module sits beside this file, not on the path of whoever runs it
sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

from handover_sections import (  # noqa: E402
    AI_MODELS,
    AI_NOTES,
    ASSISTANT_EXCLUSIONS,
    ASSISTANT_SOURCES,
    BACKGROUND,
    PAGES,
    STACK,
    UPLOAD_RULES,
    UPLOADS_TYPES,
)

DEFAULT_OUT = (
    Path.home() / "Downloads" / "Tomorrow Agent Hub - Accounts, Access and Open Items.docx"
)

# The document lists the demo logins so it is usable on its own, but this
# repository is public and passwords do not belong in it. They are supplied at
# build time; without them the document says where to ask instead.
PLACEHOLDER = "(ask the project owner)"

_parser = argparse.ArgumentParser(description="Build the handover document.")
_parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
_parser.add_argument(
    "--system-admin-password",
    default=os.environ.get("SYSTEM_ADMIN_PASSWORD", PLACEHOLDER),
)
_parser.add_argument(
    "--muni-admin-password",
    default=os.environ.get("DEMO_MUNI_ADMIN_PASSWORD", PLACEHOLDER),
)
_parser.add_argument(
    "--dept-user-password",
    default=os.environ.get("DEMO_DEPT_USER_PASSWORD", PLACEHOLDER),
)
_parser.add_argument(
    "--municipality-password",
    default=os.environ.get("MUNICIPALITY_PASSWORD", PLACEHOLDER),
    help="the shared password on the seven municipality accounts",
)
_args = _parser.parse_args()
OUT = str(_args.out)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)
    return t


def bullets(items, style_name="List Bullet"):
    for item in items:
        if isinstance(item, tuple):
            p = doc.add_paragraph(style=style_name)
            r = p.add_run(item[0])
            r.bold = True
            p.add_run(" — " + item[1])
        else:
            doc.add_paragraph(item, style=style_name)


# ---------------------------------------------------------------- title
title = doc.add_heading("Tomorrow Agent Hub", level=0)
sub = para("Production accounts, role access, and open items")
sub.runs[0].italic = True
sub.runs[0].font.size = Pt(12)
author = para("Lemar Canete  ·  ImpactMakery")
author.runs[0].font.size = Pt(10)
d = para("6 August 2026  ·  Production: https://spac-xi.vercel.app")
d.runs[0].font.size = Pt(9)
d.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

props = doc.core_properties
props.author = "Lemar Canete"
props.last_modified_by = "Lemar Canete"
props.title = "Tomorrow Agent Hub — accounts, access and open items"
props.comments = (
    "Handover document for Tomorrow Agent Hub: production accounts, role access, "
    "screens, uploads, technology and outstanding items."
)

# ---------------------------------------------------------------- accounts
h("1. Production accounts", 1)
para(
    "Three accounts cover the three roles. Passwords should be changed before the "
    "URL is shared outside the team."
)
table(
    ["Email", "Role", "Scope", "Password", "Lands on"],
    [
        [
            "info@impactmakery.com",
            "System admin",
            "All municipalities",
            _args.system_admin_password,
            "/system/stats",
        ],
        [
            "demo.admin@impactmakery.com",
            "Municipality admin",
            "עיריית דוגמה",
            _args.muni_admin_password,
            "/admin/stats",
        ],
        [
            "demo.user@impactmakery.com",
            "Department user",
            "עיריית דוגמה / רווחה",
            _args.dept_user_password,
            "/chat",
        ],
    ],
)
para("")
h("Municipality accounts", 2)
para(
    "One administrator per municipality, created on 12 August 2026 alongside their "
    "documents. Each can reach only their own municipality. All seven share the same "
    "password for now, and all have the weekly digest switched off because the "
    "addresses are placeholders — give each its own password and a real address "
    "before staff use them."
)
table(
    ["Email", "Municipality", "Contact named in the folder", "Password"],
    [
        ["deiralasad@impactmakery.com", "דיר אל אסד", "הנד", _args.municipality_password],
        ["hurfeish@impactmakery.com", "חורפיש", "אסיה", _args.municipality_password],
        ["maaleyosef@impactmakery.com", "מעלה יוסף", "לב", _args.municipality_password],
        ["nahariya@impactmakery.com", "נהריה", "אלצ׳ין", _args.municipality_password],
        ["tzfat@impactmakery.com", "צפת", "אוריאל", _args.municipality_password],
        ["kiryatshmona@impactmakery.com", "קריית שמונה", "מתניה", _args.municipality_password],
        ["shlomi@impactmakery.com", "שלומי", "אליהו", _args.municipality_password],
    ],
)

para("")
para(
    "The two demo accounts were created directly with passwords rather than by "
    "invitation, which was the only way in before email was configured. New "
    "users should now be invited normally from the Users screen.",
    italic=True,
)

# ---------------------------------------------------------------- access
doc.add_page_break()
h("2. What each role can reach", 1)
para(
    "“Sidebar” means the link appears in the navigation. Every route is also "
    "enforced on the server: a role reaching a page it should not see gets a 404, "
    "not a 403, so the existence of another municipality's data is never revealed."
)
para("")
table(
    ["Page / feature", "Route", "System admin", "Municipality admin", "Department user"],
    [
        ["Assistant (chat)", "/chat", "Yes", "Yes", "Yes"],
        ["Shared library — browse & search", "/knowledge", "Yes", "No", "No"],
        ["Shared library — upload, edit, delete", "/knowledge", "Yes", "No", "No"],
        ["Read shared-library content", "via the assistant", "Yes", "Yes", "Yes"],
        ["Own municipality's library — browse", "/knowledge", "Any municipality",
         "Own municipality", "No"],
        ["Own municipality's library — upload, edit, delete", "/knowledge",
         "Any municipality", "Own municipality", "No"],
        ["Open a document the assistant cited", "/knowledge/[id]", "Yes", "Yes", "Yes"],
        ["Shared board (global)", "/board", "Yes", "Yes", "Yes"],
        ["Publish to a board", "/board", "Yes", "Yes", "Yes"],
        ["Delete board items", "/board", "Anywhere", "Any in own municipality", "Own items only"],
        ["Municipality board", "/municipality", "No municipality",
         "Own municipality", "Own municipality"],
        ["Department area (hidden for now)", "/departments/[id]", "Any",
         "Any in own municipality", "Own departments only"],
        ["Profile & language", "/profile", "Yes", "Yes", "Yes"],
        ["Users — invite, roles, deactivate", "/admin/users",
         "via /system/users", "Own municipality", "No"],
        ["Departments — create, archive", "/admin/departments", "No", "Own municipality", "No"],
        ["Usage dashboard", "/admin/stats", "via /system/stats", "Own municipality", "No"],
        ["Municipalities", "/system/municipalities", "Yes", "No", "No"],
        ["Categories", "/system/categories", "Yes", "No", "No"],
        ["All users, all municipalities", "/system/users", "Yes", "No", "No"],
        ["Platform usage + unanswered questions", "/system/stats", "Yes", "No", "No"],
    ],
)

# ---------------------------------------------------------------- nuances
doc.add_page_break()
h("3. Behaviour that looks like a bug but is not", 1)
para(
    "Every item below is deliberate and matches the agreed scope. They are "
    "collected here because they are the things users are most likely to report "
    "as faults — knowing them in advance saves the support conversation."
)
bullets(
    [
        (
            "The assistant only answers from what you can see",
            "the permission filter is inside the search query itself, so a user can "
            "never receive an answer drawn from material they have no access to.",
        ),
        (
            "A municipality admin can read a department's files but the "
            "assistant will not cite them",
            "browsing access and assistant access are not the same. Department content "
            "is retrievable by the assistant for department members only. An admin who "
            "needs answers from a department should be added to it.",
        ),
        (
            "Unanswered questions: count vs. text",
            "a municipality admin sees how many questions went unanswered; the wording "
            "of those questions is visible to the system admin only.",
        ),
        (
            "“Not covered by the available material” is a correct answer",
            "it means nothing relevant was found that this user may see. It is preferred "
            "over an unsupported answer.",
        ),
        (
            "Archiving a department is reversible for 90 days",
            "content is hidden immediately, including from the assistant, then permanently "
            "deleted after 90 days.",
        ),
        (
            "Deactivating a user takes effect immediately",
            "it signs them out of every device. Changing someone's role or departments "
            "also forces a fresh sign-in so the new permissions apply at once.",
        ),
        (
            "Greetings are answered conversationally",
            "saying “hello” gets a normal reply rather than “not covered”, and is not "
            "recorded as an unanswered question.",
        ),
        (
            "The assistant replies in the language of the question",
            "a Hebrew question gets a Hebrew answer even when the source document is in "
            "English, and the reverse.",
        ),
    ]
)

# ---------------------------------------------------------------- pending
doc.add_page_break()
h("4. Every screen, and what it can do", 1)
para(
    "The last column lists the server endpoints each screen uses, for whoever "
    "maintains the system. Every one of them re-checks permissions: reaching a "
    "screen or an endpoint you should not see returns “not found”, never “not "
    "allowed”, so the existence of another municipality's data is never revealed."
)
para("")
for name, route, who, does, apis in PAGES:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}  ")
    r.bold = True
    r2 = p.add_run(route)
    r2.font.name = "Consolas"
    r2.font.size = Pt(9)
    r3 = p.add_run(f"   ·   {who}")
    r3.italic = True
    r3.font.size = Pt(9)
    doc.add_paragraph(does, style="List Bullet")
    ap = doc.add_paragraph()
    ar = ap.add_run(apis)
    ar.font.name = "Consolas"
    ar.font.size = Pt(8)
    ar.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()
h("5. Files you can upload", 1)
table(
    ["Format", "Extension", "Typically used for"],
    UPLOADS_TYPES,
)
para("")
bullets(UPLOAD_RULES)

doc.add_page_break()
h("6. What runs on its own", 1)
bullets(BACKGROUND)

doc.add_page_break()
h("7. Technology", 1)
para(
    "Chosen so the whole system is one database, one API and one frontend — "
    "no separate search engine, message queue or vector store to run or pay for."
)
para("")
table(["Layer", "Technology", "Why it is there"], STACK)

doc.add_page_break()
h("8. AI models in use, and what they cost", 1)
para(
    "Two separate jobs, deliberately kept on separate providers so either can be "
    "changed without touching the other."
)
para("")
table(["Job", "Provider", "Model", "Cost"], AI_MODELS)
para("")
bullets(AI_NOTES)

doc.add_page_break()
h("9. How the assistant finds an answer", 1)
para(
    "Three independent searches run over the same material and are combined, "
    "then the best twelve passages are chosen and sent to the model. Every one "
    "of them applies the permission rules in section 3 — none is a way around."
)
para("")
table(
    ["Search", "Answers", "Without it you would lose"],
    [
        [
            "Meaning (vector)",
            "passages that mean the same thing as the question, in either language",
            "any question phrased differently from the document",
        ],
        [
            "Exact words (full text)",
            "literal terms: form 4B, regulation 17.3, a tender number",
            "reference numbers, which carry almost no meaning to match on",
        ],
        [
            "Connections (knowledge graph)",
            "what is linked to what, across documents",
            "answers assembled from two documents that share a department, "
            "programme or regulation",
        ],
    ],
)
para("")
para("")
para("Where its answers come from", bold=True)
para("")
table(
    ["Source", "What is searched", "Who can be answered from it"],
    ASSISTANT_SOURCES,
)
para("")
para("What it is not given", bold=True)
bullets(ASSISTANT_EXCLUSIONS)
para("")
bullets(
    [
        (
            "Scanned documents are read",
            "a PDF that is only images — a signed letter, a stamped approval — is "
            "read with optical character recognition in Hebrew and English. Without "
            "this such files are invisible to the assistant. One of your own signed "
            "Hebrew PDFs contains six characters of machine-readable text.",
        ),
        (
            "Passages are chosen for variety, not just similarity",
            "one long document cannot fill the whole answer and crowd out the "
            "paragraph elsewhere that completes it.",
        ),
        (
            "“Text” on a board post is searchable in full",
            "the option is labelled Text because people share prompts for AI "
            "assistants, agent instructions, and plain notes worth passing on. "
            "Whatever is written there can be found both by searching the board "
            "and by asking the assistant.",
        ),
        (
            "The knowledge graph is built by pattern matching, not a model",
            "the model version finds far more in Hebrew — pattern matching has no "
            "capitalisation to lean on, so it finds entities but few relationships. "
            "It was switched off for the first bulk load because building the graph "
            "with a model meant roughly ten thousand calls against a daily free "
            "quota. Search, ranking and citations are unaffected; the graph is a "
            "third way of finding passages, not the only one. Pointing it back at a "
            "funded model and re-indexing restores the fuller version.",
        ),
    ]
)

doc.add_page_break()
h("10. Outstanding and pending items", 1)

h("10.1  Needs your action — do these first", 2)
table(
    ["Item", "Why it matters", "Effort"],
    [
        [
            "Rotate the API keys",
            "The OpenRouter, OpenAI and Cloudflare R2 keys were shared in plain text "
            "during development. Rotate them and update the Railway variables.",
            "15 min",
        ],
        [
            "Change the two demo passwords",
            "They appear in section 1 of this document and in the chat transcript.",
            "5 min",
        ],
        [
            "Enable database backups",
            "None are configured. Railway schedules them per volume under the Postgres "
            "service's Backups tab — no plan upgrade needed. Daily keeps 6 days, weekly "
            "1 month, monthly 3 months; enable all three.",
            "5 min",
        ],
    ],
)

h("10.2  Waiting on a third-party account", 2)
table(
    ["Service", "What stays broken without it"],
    [
        [
            "Custom domain for email (optional)",
            "Email is live and sending from updates.impactmakery.com. Messages will "
            "arrive from that address rather than the main company domain until the "
            "apex domain is verified with the email provider as well.",
        ],
        [
            "Custom domain + DNS",
            "The product is on a vercel.app address. Also needed before email will look "
            "legitimate to recipients.",
        ],
        [
            "Sentry",
            "No error monitoring. Production failures are currently only visible in "
            "Railway logs, which are retained for 7 days on the current plan.",
        ],
    ],
)

h("10.3  Content and data", 2)
bullets(
    [
        (
            "Seven municipalities hold their own material",
            "344 documents were loaded from the folders supplied on 12 August 2026: "
            "קריית שמונה 311, שלומי 23, דיר אל אסד 10, מעלה יוסף 1. Each sits in that "
            "municipality's own library and is readable by that municipality alone.",
        ),
        (
            "Three municipalities sent empty folders",
            "חורפיש, נהריה and צפת exist with working accounts but no documents, so the "
            "assistant will answer “not covered” for them until someone uploads.",
        ),
        (
            "25 files could not be taken",
            "6 exceed the 25 MB limit (the largest is 111 MB) and 19 are formats the "
            "library does not read — .doc, .ppt, .xlsb, .html, .avif, .zip, and one "
            ".exe. The full list was delivered as a CSV. Re-saving the older Office "
            "files as .docx/.pptx would let them in; the oversize presentations need "
            "splitting or compressing.",
        ),
        (
            "One PDF needed repairing",
            "a tender notice carried 21 stray bytes in front of its PDF header, so the "
            "content check refused it. A copy with those bytes removed was uploaded; "
            "the original file was not modified.",
        ),
        (
            "The knowledge graph is entity-only for now",
            "extraction is set to pattern matching rather than a language model, because "
            "building the graph with a model would have meant roughly ten thousand calls "
            "against a free-tier quota. Search, ranking and citations are unaffected — "
            "the graph is a third way of finding passages, not the only one. Pointing "
            "GRAPH_EXTRACTOR back at a funded model and re-running the re-index command "
            "builds the full graph, including the Hebrew relationships pattern matching "
            "cannot see.",
        ),
        (
            "Two throwaway municipalities still exist",
            "“Sample” (empty, from earlier testing) and “עיריית דוגמה” (holds the two "
            "demo accounts). Both should be removed before real use.",
        ),
        (
            "The seven municipality accounts share one password",
            "they were created directly rather than by invitation because the addresses "
            "are placeholders. Each should be given its own password, and its real "
            "address, before staff use them. The weekly digest is switched off on all "
            "seven so nothing bounces.",
        ),
        (
            "Hebrew retrieval can now be measured, and should be",
            "until this load the search had only been tested against synthetic Hebrew. "
            "There is now a real corpus behind it, so the open question — does it answer "
            "the questions staff actually ask? — is finally answerable. Worth doing "
            "before a pilot.",
        ),
    ]
)

h("10.4  Known limits of the current setup", 2)
bullets(
    [
        (
            "Database storage 5 GB",
            "roughly 300,000 document passages — comfortable for a pilot, not for years "
            "of documents across many municipalities. Uploaded files themselves do not "
            "count towards it; they live in Cloudflare R2.",
        ),
        (
            "Log history 7 days",
            "enough for immediate debugging, not for investigating something reported "
            "late.",
        ),
        (
            "Uploads are capped at 25 MB",
            "and limited to PDF, DOCX, PPTX and XLSX.",
        ),
        (
            "Uploads are not checked for viruses",
            "the shared board passes any file type between municipalities without "
            "scanning it. Adding a virus check would not narrow what people may "
            "share, and is worth doing before the platform is widely used.",
        ),
        (
            "OCR reads the first 40 pages of a scan",
            "a longer scanned document is indexed only that far, so that one large file "
            "cannot hold up the whole ingestion queue.",
        ),
    ]
)

h("10.5  Deliberately deferred", 2)
bullets(
    [
        (
            "Agentic retrieval",
            "the assistant answers from one retrieval pass. Multi-step planning — "
            "search, read, decide what to search next — would answer harder questions "
            "at the cost of several model calls per question and noticeably slower "
            "replies. Awaiting your decision.",
        ),
        (
            "Follow-up question rewriting",
            "“and what about last year?” is not yet rewritten into a standalone query.",
        ),
        (
            "A cross-encoder re-ranker",
            "would score better than the current re-ranking, but adds roughly 2 GB to a "
            "205 MB image. Worth revisiting once real documents show where retrieval is "
            "actually weak.",
        ),
        (
            "A remove option for the demo data script",
            "deleting seeded demo content by hand means ordered deletes across about ten "
            "tables plus the stored files. Ask before doing it manually.",
        ),
    ]
)

para("")
para(
    "Everything in sections 1 to 9 is built, tested and live. Section 10 is what "
    "stands between the current state and a pilot with real users.",
    italic=True,
)

doc.save(OUT)
print("wrote:", OUT)
